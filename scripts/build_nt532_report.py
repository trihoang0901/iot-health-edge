from __future__ import annotations

import argparse
from dataclasses import dataclass
from hashlib import sha256
import json
import math
from pathlib import Path
import re
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "deliverables" / "BAO-CAO-NT532-MQTT-MVP.md"
DEFAULT_OUTPUT = ROOT / "deliverables" / "BAO-CAO-NT532-MQTT-MVP.docx"
ASSET_DIR = ROOT / ".codex-tmp" / "report-assets"
COVER_FRAME = ROOT / "deliverables" / "assets" / "uit-cover-frame.jpg"
COVER_LOGO = ROOT / "deliverables" / "assets" / "uit-logo.jpg"
RQ1_SOURCE_FILES = (
    "edge/db.py",
    "edge/rules.py",
    "edge/service.py",
    "edge/mqtt_client.py",
)
BASELINE_COMMIT = "7030e4b30300dec65646e3091356ca00d9eaa8f5"
BASELINE_RQ1_SOURCE_SHA256 = "760429f9dceed614279cb6c937d111a66fb1cb63ca813ed615c7de1bbd24c280"
UI_SOURCE_FILES = (
    "edge/static/app.js",
    "edge/static/index.html",
    "edge/static/styles.css",
    "scripts/dashboard-browser-smoke.js",
)

PAGE_WIDTH_DXA = 9360
LANDSCAPE_WIDTH_DXA = 12960
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

NAVY = "18324A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6770"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GOLD = "9A6B00"
GOLD_FILL = "FFF4D6"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "111111"


@dataclass(frozen=True)
class NumberingIds:
    bullet: int
    decimal: int


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, **kwargs: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ("top", "start", "bottom", "end"):
        if margin not in kwargs:
            continue
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(kwargs[margin]))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def shade_paragraph(paragraph, fill: str, *, left_border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    if left_border:
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), left_border)
        borders.append(left)


def box_paragraph(paragraph, *, color: str = BLACK, size: int = 12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "5")
        edge.set(qn("w:color"), color)


def make_picture_floating_behind(
    shape,
    *,
    x: int,
    y: int,
    alt_text: str,
) -> None:
    """Convert a python-docx inline picture to a page-relative background anchor."""
    anchor = shape._inline
    anchor.tag = qn("wp:anchor")
    for name, value in {
        "distT": "0",
        "distB": "0",
        "distL": "0",
        "distR": "0",
        "simplePos": "0",
        "relativeHeight": "251658240",
        "behindDoc": "1",
        "locked": "0",
        "layoutInCell": "1",
        "allowOverlap": "1",
    }.items():
        anchor.set(name, value)

    simple_pos = OxmlElement("wp:simplePos")
    simple_pos.set("x", "0")
    simple_pos.set("y", "0")

    position_h = OxmlElement("wp:positionH")
    position_h.set("relativeFrom", "page")
    pos_h = OxmlElement("wp:posOffset")
    pos_h.text = str(x)
    position_h.append(pos_h)

    position_v = OxmlElement("wp:positionV")
    position_v.set("relativeFrom", "page")
    pos_v = OxmlElement("wp:posOffset")
    pos_v.text = str(y)
    position_v.append(pos_v)

    effect_extent = OxmlElement("wp:effectExtent")
    for edge in ("l", "t", "r", "b"):
        effect_extent.set(edge, "0")
    wrap_none = OxmlElement("wp:wrapNone")

    anchor.insert(0, simple_pos)
    anchor.insert(1, position_h)
    anchor.insert(2, position_v)
    anchor.insert(4, effect_extent)
    anchor.insert(5, wrap_none)
    anchor.docPr.set("name", "UIT ornamental cover frame")
    anchor.docPr.set("descr", alt_text)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Mục lục sẽ được Word/LibreOffice cập nhật khi mở tài liệu."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\)|\*[^*]+\*)"
)


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor : match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), bold=True)
        elif token.startswith("`"):
            set_run_font(
                paragraph.add_run(token[1:-1]),
                name="Consolas",
                size=9,
                color=DARK_BLUE,
            )
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()  # type: ignore[union-attr]
            add_hyperlink(paragraph, label, url)
        elif token.startswith("*"):
            set_run_font(paragraph.add_run(token[1:-1]), italic=True)
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]))


def add_numbering_definition(document: Document, *, kind: str) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "540")
    indentation.set(qn("w:hanging"), "279")
    p_pr.extend([tabs, indentation])
    level.extend([start, number_format, level_text, justification, p_pr])
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref)
    numbering.append(number)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_node])
    p_pr.append(num_pr)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Code Block" not in [style.name for style in document.styles]:
        code_style = document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = document.styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(8.5)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.10)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.line_spacing = 1.0


def configure_section(section, *, landscape: bool = False) -> None:
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    section.top_margin = Inches(0.98)
    section.right_margin = Inches(0.79)
    section.bottom_margin = Inches(0.95)
    section.left_margin = Inches(0.98)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section, *, first: bool = False) -> None:
    if first:
        section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.text = ""
    run = p.add_run("NT532 · IoT Protocol · MQTT Edge Reliability")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    if first:
        first_header = section.first_page_header
        first_header.is_linked_to_previous = False
        first_header.paragraphs[0].text = ""

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    run = fp.add_run("NT532  |  Trang ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(fp)
    if first:
        first_footer = section.first_page_footer
        first_footer.is_linked_to_previous = False
        first_footer.paragraphs[0].text = ""


def add_cover(document: Document, metadata: dict[str, str]) -> None:
    required = (
        "institution",
        "faculty",
        "lecturer",
        "class_code",
        "group",
        "student_1_name",
        "student_1_id",
        "student_2_name",
        "student_2_id",
        "semester",
        "submission_place",
        "submission_date",
    )
    missing = [
        key
        for key in required
        if not metadata.get(key) or "CHƯA CUNG CẤP" in metadata.get(key, "")
    ]
    missing_assets = [str(path) for path in (COVER_FRAME, COVER_LOGO) if not path.is_file()]
    if missing or missing_assets:
        raise ValueError(
            "Cannot build submission cover; "
            f"missing metadata={missing}, missing assets={missing_assets}"
        )

    date_parts = metadata["submission_date"].split("/")
    if len(date_parts) != 3:
        raise ValueError("submission_date must use DD/MM/YYYY")
    day, month, year = date_parts

    frame_paragraph = document.add_paragraph()
    frame_paragraph.paragraph_format.space_before = Pt(0)
    frame_paragraph.paragraph_format.space_after = Pt(0)
    frame_paragraph.paragraph_format.line_spacing = Pt(1)
    frame_shape = frame_paragraph.add_run().add_picture(
        str(COVER_FRAME),
        width=Inches(7.24),
        height=Inches(10.41),
    )
    make_picture_floating_behind(
        frame_shape,
        x=Inches(0.47),
        y=Inches(0.44),
        alt_text="Khung trang trí bìa theo mẫu báo cáo UIT do người dùng cung cấp.",
    )

    identity_lines = (
        "ĐẠI HỌC QUỐC GIA THÀNH PHỐ HỒ CHÍ MINH",
        "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN",
        metadata["faculty"].upper(),
    )
    for index, text in enumerate(identity_lines):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0 if index < 2 else 3)
        set_run_font(
            p.add_run(text),
            name="Times New Roman",
            size=12,
            color=BLACK,
            bold=True,
        )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(5)
    logo = p.add_run().add_picture(str(COVER_LOGO), width=Inches(1.58))
    logo._inline.docPr.set("name", "Logo UIT")
    logo._inline.docPr.set(
        "descr", "Logo Trường Đại học Công nghệ Thông tin, ĐHQG-HCM."
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    set_run_font(
        p.add_run("BÁO CÁO ĐỒ ÁN"),
        name="Times New Roman",
        size=18,
        color=BLACK,
        bold=True,
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run_font(
        p.add_run("MÔN HỌC: CÔNG NGHỆ INTERNET OF THINGS HIỆN ĐẠI (NT532)"),
        name="Times New Roman",
        size=12,
        color=BLACK,
        bold=True,
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_run_font(
        p.add_run("ĐỀ TÀI"),
        name="Times New Roman",
        size=12,
        color=BLACK,
        bold=True,
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    set_run_font(
        p.add_run(
            "ĐÁNH GIÁ ĐỘ TIN CẬY XỬ LÝ BẢN TIN MQTT\n"
            "TRONG HỆ THỐNG IOT EDGE PHI LÂM SÀNG\n"
            "DƯỚI LỖI CẢM BIẾN VÀ LỖI Ở TẦNG ỨNG DỤNG"
        ),
        name="Times New Roman",
        size=15,
        color=BLACK,
        bold=True,
    )

    for label, value in (
        ("Giảng viên hướng dẫn", metadata["lecturer"]),
        ("Mã lớp", metadata["class_code"]),
        ("Nhóm", metadata["group"]),
    ):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(
            p.add_run(f"{label}: "),
            name="Times New Roman",
            size=11.5,
            color=BLACK,
            bold=True,
        )
        set_run_font(
            p.add_run(value),
            name="Times New Roman",
            size=11.5,
            color=BLACK,
        )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(1.05)
    p.paragraph_format.right_indent = Inches(1.05)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_together = True
    box_paragraph(p, color=BLACK, size=10)
    set_run_font(
        p.add_run("SINH VIÊN THỰC HIỆN\n"),
        name="Times New Roman",
        size=11.5,
        color=BLACK,
        bold=True,
    )
    set_run_font(
        p.add_run(f"{metadata['student_1_id']} – {metadata['student_1_name']}\n"),
        name="Times New Roman",
        size=11.5,
        color=BLACK,
    )
    set_run_font(
        p.add_run(f"{metadata['student_2_id']} – {metadata['student_2_name']}"),
        name="Times New Roman",
        size=11.5,
        color=BLACK,
    )

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    set_run_font(
        p.add_run(f"HỌC KỲ {metadata['semester'].upper()}"),
        name="Times New Roman",
        size=11,
        color=BLACK,
        bold=True,
    )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    set_run_font(
        p.add_run(
            f"{metadata['submission_place']}, ngày {day} tháng {month} năm {year}"
        ),
        name="Times New Roman",
        size=11,
        color=BLACK,
    )
    document.add_page_break()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines) or "|" not in lines[start]:
        return None
    separator = lines[start + 1].strip()
    if not re.match(r"^\|?\s*:?-{3,}", separator):
        return None
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and "|" in lines[index] and lines[index].strip():
        raw = lines[index].strip().strip("|")
        rows.append([cell.strip() for cell in raw.split("|")])
        index += 1
    rows.pop(1)
    return rows, index


def compute_widths(rows: list[list[str]], total_width: int) -> list[int]:
    columns = len(rows[0])
    maxima: list[int] = []
    for column in range(columns):
        maximum = max(len(re.sub(r"[`*_\[\]]", "", row[column])) for row in rows)
        maxima.append(maximum)
    weights = [max(0.8, min(4.5, math.sqrt(value + 4))) for value in maxima]
    raw = [max(720, int(total_width * weight / sum(weights))) for weight in weights]
    difference = total_width - sum(raw)
    raw[-1] += difference
    if raw[-1] < 720:
        deficit = 720 - raw[-1]
        raw[-1] = 720
        widest = max(range(len(raw) - 1), key=raw.__getitem__)
        raw[widest] -= deficit
    return raw


def set_table_geometry(table, widths: list[int], total_width: int) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, **CELL_MARGINS_DXA)


def add_markdown_table(document: Document, rows: list[list[str]], *, landscape: bool) -> None:
    total_width = LANDSCAPE_WIDTH_DXA if landscape else PAGE_WIDTH_DXA
    widths = compute_widths(rows, total_width)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_geometry(table, widths, total_width)
    # Six-column traceability tables need a slightly tighter landscape token to
    # keep the final row on the same page without relying on Word autofit.
    compact_font = 6.5 if landscape else 8.6
    short_columns = [
        max(len(row[column]) for row in rows) <= 18 for column in range(len(rows[0]))
    ]
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0 if landscape else 2)
            paragraph.paragraph_format.line_spacing = 1.0 if landscape else 1.08
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if short_columns[column_index]
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            add_inline(paragraph, value)
            for run in paragraph.runs:
                set_run_font(
                    run,
                    size=compact_font,
                    color=WHITE if row_index == 0 else BLACK,
                    bold=True if row_index == 0 else None,
                )
            if row_index == 0:
                shade_cell(cell, NAVY)
            elif row_index % 2 == 0:
                shade_cell(cell, LIGHT_GRAY)
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_tr_pr.append(repeat)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_picture(document: Document, path: Path, caption: str, alt_text: str, *, width: float = 6.15) -> None:
    if not path.is_file():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    cp = document.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(2)
    cp.paragraph_format.space_after = Pt(8)
    cp.paragraph_format.keep_together = True
    set_run_font(cp.add_run(caption), size=9, color=MUTED, italic=True)


def load_font(size: int, *, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.is_file():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, *, width: int, font, fill: str, spacing: int = 8) -> int:
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        candidate = f"{line} {word}".strip()
        if draw.textlength(candidate, font=font) <= width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    x, y = xy
    for entry in lines:
        draw.text((x, y), entry, font=font, fill=fill)
        y += font.size + spacing
    return y


def generate_rq1_figure() -> Path:
    baseline_path = ROOT / "evidence" / "analysis" / "baseline-reliability.json"
    hardened_path = ROOT / "evidence" / "analysis" / "hardened-reliability.json"
    if not baseline_path.is_file() or not hardened_path.is_file():
        raise FileNotFoundError("missing required RQ1 evidence artifacts")
    baseline_payload = json.loads(baseline_path.read_text(encoding="utf-8"))
    hardened_payload = json.loads(hardened_path.read_text(encoding="utf-8"))
    digest = sha256()
    for relative in RQ1_SOURCE_FILES:
        digest.update(relative.encode("utf-8"))
        digest.update(b"\0")
        digest.update((ROOT / relative).read_bytes())
        digest.update(b"\0")
    baseline_provenance = baseline_payload.get("source_provenance", {})
    hardened_provenance = hardened_payload.get("source_provenance", {})
    for payload, implementation in (
        (baseline_payload, "baseline"),
        (hardened_payload, "hardened"),
    ):
        if payload.get("artifact_version") != "1.0":
            raise ValueError(f"RQ1 {implementation} artifact version mismatch")
        if payload.get("implementation") != implementation:
            raise ValueError(f"RQ1 {implementation} label mismatch")
        if payload.get("repetitions") != 30 or len(payload.get("runs", [])) != 30:
            raise ValueError(f"RQ1 {implementation} must contain 30 runs")
        if payload.get("deterministic_repeatability_only") is not True:
            raise ValueError(f"RQ1 {implementation} repeatability boundary missing")
        if payload.get("inferential_confidence_interval") is not None:
            raise ValueError(f"RQ1 {implementation} must not claim inferential CI")
        runs = payload["runs"]
        if [run.get("repetition") for run in runs] != list(range(1, 31)):
            raise ValueError(f"RQ1 {implementation} repetition identity mismatch")
        recomputed_cases = {
            "atomic_alert_passed": sum(
                run.get("atomic_alert", {}).get("pass") is True for run in runs
            ),
            "old_lwt_session_passed": sum(
                run.get("old_lwt_session", {}).get("pass") is True for run in runs
            ),
        }
        if payload.get("cases") != recomputed_cases:
            raise ValueError(f"RQ1 {implementation} case reconciliation failed")
    if baseline_provenance.get("source_state") != "commit_clean":
        raise ValueError("RQ1 baseline source must be clean")
    if baseline_provenance.get("rq1_source_sha256") != BASELINE_RQ1_SOURCE_SHA256:
        raise ValueError("RQ1 baseline scoped source hash mismatch")
    if hardened_provenance.get("rq1_source_files") != list(RQ1_SOURCE_FILES):
        raise ValueError("RQ1 hardened source file scope mismatch")
    if hardened_provenance.get("rq1_source_sha256") != digest.hexdigest():
        raise ValueError("RQ1 hardened evidence is stale for the current source")
    if (
        baseline_payload.get("commit") != BASELINE_COMMIT
        or baseline_payload.get("baseline_commit") != BASELINE_COMMIT
        or hardened_payload.get("baseline_commit") != BASELINE_COMMIT
    ):
        raise ValueError("RQ1 baseline evidence is not pinned to its baseline commit")
    if baseline_payload.get("cases") != {
        "atomic_alert_passed": 0,
        "old_lwt_session_passed": 0,
    }:
        raise ValueError("RQ1 baseline counterexamples were not reproduced")
    if hardened_payload.get("cases") != {
        "atomic_alert_passed": 30,
        "old_lwt_session_passed": 30,
    }:
        raise ValueError("RQ1 hardened repeatability did not reach 30/30")
    if any(
        run.get("old_lwt_session", {}).get("stale_disposition") != "stale"
        for run in hardened_payload["runs"]
    ):
        raise ValueError("RQ1 hardened old LWT disposition must be stale")
    baseline = baseline_payload["cases"]
    hardened = hardened_payload["cases"]
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    output = ASSET_DIR / "rq1-repeatability.png"
    image = Image.new("RGB", (1600, 860), "#F7F9FC")
    draw = ImageDraw.Draw(image)
    title_font = load_font(52, bold=True)
    label_font = load_font(34, bold=True)
    value_font = load_font(44, bold=True)
    note_font = load_font(25)
    draw.text((90, 55), "RQ1 · Repeatability dưới fault fixture", font=title_font, fill=f"#{NAVY}")
    draw.text((90, 125), "30 lần lặp deterministic · không dùng CI suy diễn", font=note_font, fill=f"#{MUTED}")
    cases = [
        ("Atomic telemetry → rule → alert", baseline["atomic_alert_passed"], hardened["atomic_alert_passed"]),
        ("Old LWT không rewind session", baseline["old_lwt_session_passed"], hardened["old_lwt_session_passed"]),
    ]
    x0, chart_y, bar_width, gap = 635, 235, 760, 58
    for index, (label, base, hard) in enumerate(cases):
        y = chart_y + index * 280
        draw_wrapped(draw, (90, y + 15), label, width=470, font=label_font, fill=f"#{BLACK}")
        for row, (name, value, color) in enumerate(
            (("Baseline", base, "#C9675A"), ("Hardened", hard, "#2E8064"))
        ):
            by = y + row * 92
            draw.rounded_rectangle((x0, by, x0 + bar_width, by + 60), radius=18, fill="#E1E6EC")
            filled = max(8, int(bar_width * value / 30)) if value else 0
            if filled:
                draw.rounded_rectangle((x0, by, x0 + filled, by + 60), radius=18, fill=color)
            draw.text((x0 - 155, by + 11), name, font=note_font, fill=f"#{MUTED}")
            draw.text((x0 + bar_width + 25, by + 3), f"{value}/30", font=value_font, fill=color)
    draw.text((90, 790), "Nguồn: evidence/analysis/{baseline,hardened}-reliability.json", font=note_font, fill=f"#{MUTED}")
    image.save(output, optimize=True)
    return output


def generate_rq2_figure() -> Path:
    source = ROOT / "evidence" / "analysis" / "rq2-v5-experiments.json"
    if not source.is_file():
        raise FileNotFoundError("missing required RQ2 v5 aggregate")
    data = json.loads(source.read_text(encoding="utf-8"))
    if data.get("artifact_version") != "5.0":
        raise ValueError("RQ2 aggregate artifact_version must be 5.0")
    if data.get("run_prefix") != "nt532-rq2-v5-":
        raise ValueError("RQ2 aggregate run_prefix must be nt532-rq2-v5-")
    if data.get("matched_seed_count") != 30:
        raise ValueError("RQ2 aggregate must contain exactly 30 matched seeds")
    profiles = data["profiles"]
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    output = ASSET_DIR / "rq2-latency.png"
    image = Image.new("RGB", (1600, 980), "#F7F9FC")
    draw = ImageDraw.Draw(image)
    title_font = load_font(52, bold=True)
    label_font = load_font(31, bold=True)
    value_font = load_font(42, bold=True)
    note_font = load_font(25)
    draw.text(
        (90, 55),
        "RQ2 · Coverage lịch phát và độ trễ upper-bound",
        font=title_font,
        fill=f"#{NAVY}",
    )
    draw.text(
        (90, 125),
        "Median theo run · 30 seed độc lập / profile · cùng host monotonic",
        font=note_font,
        fill=f"#{MUTED}",
    )
    latency_values = [
        float(profiles[profile]["metrics"][metric]["median"])
        for profile in ("lan-baseline", "remote-app-emulated")
        for metric in (
            "schedule_to_api_upper_bound_p50_ms",
            "schedule_to_api_upper_bound_p95_ms",
        )
    ]
    max_value = max(1.0, max(latency_values) * 1.12)
    x0, y0, width = 610, 230, 650
    colors = {"lan-baseline": "#2E74B5", "remote-app-emulated": "#C18A1B"}
    for p_index, profile in enumerate(("lan-baseline", "remote-app-emulated")):
        metrics = profiles[profile]["metrics"]
        label = "LAN baseline" if profile == "lan-baseline" else "Nhiễu tầng ứng dụng"
        block_y = y0 + p_index * 325
        draw.text((90, block_y), label, font=label_font, fill=f"#{BLACK}")
        coverage = float(metrics["scheduled_observation_ratio"]["median"])
        coverage_y = block_y + 62
        draw.text((295, coverage_y + 6), "coverage", font=note_font, fill=f"#{MUTED}")
        draw.rounded_rectangle(
            (x0, coverage_y, x0 + width, coverage_y + 48),
            radius=15,
            fill="#E1E6EC",
        )
        coverage_width = max(0, min(width, int(width * coverage)))
        if coverage_width:
            draw.rounded_rectangle(
                (x0, coverage_y, x0 + coverage_width, coverage_y + 48),
                radius=15,
                fill=colors[profile],
            )
        draw.text(
            (x0 + width + 24, coverage_y - 2),
            f"{coverage * 100:.1f}%",
            font=value_font,
            fill=colors[profile],
        )
        for m_index, (metric, short) in enumerate(
            (
                ("schedule_to_api_upper_bound_p50_ms", "p50"),
                ("schedule_to_api_upper_bound_p95_ms", "p95"),
            )
        ):
            value = float(metrics[metric]["median"])
            y = block_y + 130 + m_index * 76
            draw.text((390, y + 5), f"schedule→API {short}", font=note_font, fill=f"#{MUTED}")
            draw.rounded_rectangle((x0, y, x0 + width, y + 50), radius=16, fill="#E1E6EC")
            filled = int(width * value / max_value)
            draw.rounded_rectangle((x0, y, x0 + filled, y + 50), radius=16, fill=colors[profile])
            draw.text((x0 + width + 25, y), f"{value:.2f} ms", font=value_font, fill=colors[profile])
    draw.text(
        (90, 880),
        "Coverage dùng toàn bộ lịch phát; latency chỉ tính mẫu được API quan sát.",
        font=note_font,
        fill=f"#{BLACK}",
    )
    draw.text(
        (90, 920),
        "App impairment trước publish; network_claim=none; measured_5g=false.",
        font=note_font,
        fill=f"#{RED}",
    )
    image.save(output, optimize=True)
    return output


def generate_dashboard_crop() -> Path:
    source = ROOT / "evidence" / "ui" / "dashboard-desktop-1440.png"
    report_path = ROOT / "evidence" / "ui" / "browser-smoke.json"
    if not source.is_file():
        raise FileNotFoundError("missing required dashboard browser evidence")
    if not report_path.is_file():
        raise FileNotFoundError("missing required browser smoke report")
    report = json.loads(report_path.read_text(encoding="utf-8"))
    if report.get("artifact_version") != "1.1" or report.get("status") != "passed":
        raise ValueError("browser smoke report is not a completed v1.1 pass")
    checks = report.get("checks")
    expected_names = {"mobile-320", "mobile-360", "tablet-768", "desktop-1440"}
    if not isinstance(checks, list) or {item.get("name") for item in checks} != expected_names:
        raise ValueError("browser smoke must contain all four canonical viewports")
    provenance = report.get("source_provenance", {})
    digest = sha256()
    for relative in UI_SOURCE_FILES:
        digest.update(relative.encode("utf-8"))
        digest.update(b"\0")
        digest.update((ROOT / relative).read_bytes())
        digest.update(b"\0")
    if provenance.get("source_files") != list(UI_SOURCE_FILES):
        raise ValueError("browser evidence source scope mismatch")
    if provenance.get("source_sha256") != digest.hexdigest():
        raise ValueError("browser evidence is stale for current dashboard source")
    for item in checks:
        screenshot_name = item.get("screenshot")
        screenshot_path = ROOT / "evidence" / "ui" / str(screenshot_name)
        if screenshot_path.name != screenshot_name or not screenshot_path.is_file():
            raise ValueError("browser evidence screenshot identity mismatch")
        if item.get("screenshot_sha256") != sha256(screenshot_path.read_bytes()).hexdigest():
            raise ValueError("browser evidence screenshot hash mismatch")
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    output = ASSET_DIR / "dashboard-cockpit.png"
    with Image.open(source) as image:
        crop_height = min(image.height, int(image.width * 0.93))
        cropped = image.crop((0, 0, image.width, crop_height)).convert("RGB")
        cropped.save(output, quality=92, optimize=True)
    return output


def add_body_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    add_inline(p, text)


def parse_frontmatter(lines: list[str]) -> tuple[dict[str, str], list[str]]:
    if not lines or lines[0].strip() != "---":
        return {}, lines
    metadata: dict[str, str] = {}
    for index in range(1, len(lines)):
        if lines[index].strip() == "---":
            return metadata, lines[index + 1 :]
        if ":" in lines[index]:
            key, value = lines[index].split(":", 1)
            metadata[key.strip()] = value.strip().strip('"')
    return metadata, lines


def build_document(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    metadata, lines = parse_frontmatter(lines)
    start = next(
        (index for index, line in enumerate(lines) if line.startswith("## Quy ước")),
        0,
    )
    lines = lines[start:]

    document = Document()
    document.core_properties.title = metadata.get("title", "Báo cáo NT532 MQTT MVP")
    document.core_properties.subject = "Công nghệ IoT hiện đại NT532 · IoT Protocol"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.keywords = "MQTT 3.1.1, IoT Protocol, edge, reliability"
    configure_styles(document)
    configure_section(document.sections[0])
    configure_header_footer(document.sections[0], first=True)
    add_cover(document, metadata)

    toc_heading = document.add_paragraph("Mục lục", style="Heading 1")
    toc_heading.paragraph_format.page_break_before = False
    toc = document.add_paragraph()
    add_toc(toc)
    document.add_page_break()

    numbering = NumberingIds(
        bullet=add_numbering_definition(document, kind="bullet"),
        decimal=add_numbering_definition(document, kind="decimal"),
    )
    active_decimal_id: int | None = None
    rq1_figure = generate_rq1_figure()
    rq2_figure = generate_rq2_figure()
    dashboard_figure = generate_dashboard_crop()

    index = 0
    pending_figure: tuple[Path | None, str, str] | None = None
    landscape_section_open = False
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        is_decimal_line = bool(re.match(r"^\s*\d+\.\s+", line))
        if not is_decimal_line:
            active_decimal_id = None
        if not stripped:
            index += 1
            continue
        table_result = parse_table(lines, index)
        if table_result is not None:
            rows, next_index = table_result
            landscape = len(rows[0]) >= 6
            if landscape and not landscape_section_open:
                section = document.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=True)
                configure_header_footer(section)
                landscape_section_open = True
            add_markdown_table(document, rows, landscape=landscape)
            if pending_figure and pending_figure[0]:
                add_picture(document, pending_figure[0], pending_figure[1], pending_figure[2])
                pending_figure = None
            if landscape:
                section = document.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=False)
                configure_header_footer(section)
                landscape_section_open = False
            index = next_index
            continue
        if stripped.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            p = document.add_paragraph(style="Code Block")
            shade_paragraph(p, LIGHT_GRAY, left_border=DARK_BLUE)
            run = p.add_run("\n".join(code_lines))
            set_run_font(run, name="Consolas", size=8.5, color=BLACK)
            index += 1
            continue
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            lookahead = index + 1
            while lookahead < len(lines) and not lines[lookahead].strip():
                lookahead += 1
            next_table = parse_table(lines, lookahead) if lookahead < len(lines) else None
            if next_table is not None and len(next_table[0][0]) >= 6:
                section = document.add_section(WD_SECTION.NEW_PAGE)
                configure_section(section, landscape=True)
                configure_header_footer(section)
                landscape_section_open = True
            p = document.add_paragraph(text, style=f"Heading {level}")
            numbered_heading = re.match(r"^(\d+)\.", text)
            if not landscape_section_open and level == 1 and (
                (numbered_heading and numbered_heading.group(1) != "1")
                or text.startswith("Tài liệu tham khảo")
                or text.startswith("Phụ lục")
            ):
                p.paragraph_format.page_break_before = True
            if "6.3. Experiment cockpit" in text and dashboard_figure:
                add_picture(
                    document,
                    dashboard_figure,
                    "Hình 6.1. Cockpit MQTT trên Docker live, tách Edge, node và evidence run.",
                    "Dashboard tối với trạng thái Edge, node cảm biến, evidence run và KPI độ tin cậy MQTT.",
                )
            if "7.2." in text:
                pending_figure = (
                    rq1_figure,
                    "Hình 7.1. Repeatability RQ1 trên hai counterexample baseline và bản hardened.",
                    "Biểu đồ cho thấy baseline đạt 0 trên 30, bản hardened đạt 30 trên 30 cho atomic alert và old LWT session.",
                )
            elif "7.3." in text:
                pending_figure = (
                    rq2_figure,
                    "Hình 7.2. Coverage lịch phát và schedule-to-API polling upper-bound của hai profile RQ2.",
                    "Biểu đồ tỷ lệ API quan sát trên toàn bộ lịch phát cùng p50 và p95 schedule-to-API polling upper-bound cho LAN baseline và nhiễu tầng ứng dụng.",
                )
            index += 1
            continue
        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip().lstrip(">").strip())
                index += 1
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            shade_paragraph(p, GOLD_FILL, left_border=GOLD)
            add_inline(p, " ".join(quote_lines))
            continue
        list_match = re.match(r"^\s*([-*])\s+(.+)$", line)
        number_match = re.match(r"^\s*(\d+)\.\s+(.+)$", line)
        if list_match or number_match:
            text = (list_match or number_match).group(2)  # type: ignore[union-attr]
            index += 1
            while index < len(lines):
                continuation = lines[index]
                if not continuation.strip():
                    break
                if not continuation[:1].isspace():
                    break
                if re.match(r"^\s*([-*]|\d+\.)\s+", continuation):
                    break
                text += " " + continuation.strip()
                index += 1
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            if list_match:
                list_num_id = numbering.bullet
            else:
                if active_decimal_id is None:
                    active_decimal_id = add_numbering_definition(document, kind="decimal")
                list_num_id = active_decimal_id
            apply_numbering(p, list_num_id)
            add_inline(p, text)
            continue
        if stripped in {"---", "***"}:
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if (
                candidate.startswith(("#", ">", "```"))
                or re.match(r"^[-*]\s+", candidate)
                or re.match(r"^\d+\.\s+", candidate)
                or parse_table(lines, index) is not None
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        add_body_paragraph(document, " ".join(paragraph_lines))

    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the NT532 academic report DOCX.")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    build_document(args.source.resolve(), args.output.resolve())
    print(json.dumps({"output": str(args.output.resolve()), "preset": "narrative_proposal"}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
